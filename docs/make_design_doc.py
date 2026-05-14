# -*- coding: utf-8 -*-
"""
売上計上システム設計書 DOCX 生成スクリプト
使い方: python make_design_doc.py
出力  : 売上計上システム設計書.docx（同フォルダ内）
依存  : pip install python-docx
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

FONT = 'メイリオ'
C_NAVY  = RGBColor(0x1e, 0x3a, 0x5f)
C_BLUE  = RGBColor(0x25, 0x63, 0xeb)
C_GREEN = RGBColor(0x16, 0xa3, 0x4a)
C_ORANGE= RGBColor(0xea, 0x58, 0x0c)
C_GRAY  = RGBColor(0x64, 0x74, 0x8b)
C_WHITE = RGBColor(0xff, 0xff, 0xff)
C_HDR   = RGBColor(0x1e, 0x3a, 0x5f)

def bg(rgb): return f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}'

def cell_bg(cell, rgb):
    tc=cell._tc; pr=tc.get_or_add_tcPr(); s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),bg(rgb)); pr.append(s)

def cell_bd(cell, color='CCCCCC', sz=4):
    tc=cell._tc; pr=tc.get_or_add_tcPr(); bd=OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        e=OxmlElement(f'w:{side}'); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),str(sz)); e.set(qn('w:color'),color); bd.append(e)
    pr.append(bd)

def run(para, text, pt=10.5, bold=False, color=None):
    r=para.add_run(text); r.font.name=FONT; r.font.size=Pt(pt); r.font.bold=bold
    if color: r.font.color.rgb=color
    rf=OxmlElement('w:rFonts'); rf.set(qn('w:eastAsia'),FONT); r._element.get_or_add_rPr().insert(0,rf)
    return r

def para(text='', pt=10.5, bold=False, color=None, align=None, sb=0, sa=6):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
    if align: p.alignment=align
    if text: run(p,text,pt=pt,bold=bold,color=color)
    return p

def h1(text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(16); p.paragraph_format.space_after=Pt(6)
    pr=p._p.get_or_add_pPr(); bd=OxmlElement('w:pBdr'); l=OxmlElement('w:left')
    l.set(qn('w:val'),'single'); l.set(qn('w:sz'),'24'); l.set(qn('w:color'),'1e3a5f'); l.set(qn('w:space'),'8')
    bd.append(l); pr.append(bd); run(p,text,pt=14,bold=True,color=C_NAVY)

def h2(text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4)
    run(p,'▌ '+text,pt=11.5,bold=True,color=C_BLUE)

def table(headers, rows, widths):
    t=doc.add_table(rows=1+len(rows),cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    hr=t.rows[0]
    for i,(h,w) in enumerate(zip(headers,widths)):
        c=hr.cells[i]; c.width=w; cell_bg(c,C_HDR); cell_bd(c,'FFFFFF',6)
        p=c.paragraphs[0]; p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
        run(p,h,pt=9.5,bold=True,color=C_WHITE)
    for ri,row in enumerate(rows):
        tr=t.rows[ri+1]; rbg=RGBColor(0xf8,0xfa,0xfc) if ri%2==0 else C_WHITE
        for ci,(v,w) in enumerate(zip(row,widths)):
            c=tr.cells[ci]; c.width=w; cell_bg(c,rbg); cell_bd(c,'D1D5DB',4)
            p=c.paragraphs[0]; p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
            run(p,str(v),pt=9.5,bold=(ci==0))

def note(text, bc='93C5FD'):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(8)
    pr=p._p.get_or_add_pPr(); bd=OxmlElement('w:pBdr')
    for side in ['top','left','bottom','right']:
        e=OxmlElement(f'w:{side}'); e.set(qn('w:val'),'single')
        e.set(qn('w:sz'),'8' if side=='left' else '4'); e.set(qn('w:color'),bc); e.set(qn('w:space'),'4'); bd.append(e)
    pr.append(bd); run(p,text,pt=9.5,color=RGBColor(0x1e,0x40,0xaf))

# ═══ 表紙 ═══
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(60); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
run(p,'配車管理システム',pt=20,bold=True,color=C_NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
run(p,'売上計上モジュール 設計書',pt=26,bold=True,color=C_NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(8)
run(p,'作業指示書 × T-PLANNER 連携設計',pt=14,color=C_BLUE)
doc.add_paragraph(); doc.add_paragraph()
it=doc.add_table(rows=3,cols=2); it.alignment=WD_TABLE_ALIGNMENT.CENTER
for ri,(k,v) in enumerate([('作成日','2026年5月14日'),('バージョン','1.0（設計案）'),('対象システム','廃棄物収集・運搬 配車管理システム')]):
    r=it.rows[ri]; c0,c1=r.cells[0],r.cells[1]; c0.width=Cm(4); c1.width=Cm(8)
    cell_bg(c0,C_HDR); cell_bg(c1,RGBColor(0xf0,0xf4,0xf8)); cell_bd(c0,'1e3a5f'); cell_bd(c1,'E2E8F0')
    run(c0.paragraphs[0],k,pt=10,bold=True,color=C_WHITE); run(c1.paragraphs[0],v,pt=10)
doc.add_page_break()

# ═══ 第1章 ═══
h1('第1章　システム全体概要')
para('本システムは配車管理システムで管理している配車案件を起点として、作業指示書を発行し、T-PLANNERへの売上伝票計上まで一貫して管理するモジュールです。',pt=10.5,sa=8)
h2('1-1　全体データフロー')
flow=[('①','配車カード作成','配車管理システム上で案件を登録'),('②','売上計上区分選択','計上あり / 計上なし / 保留 を選択'),
      ('③','作業指示書番号 自動採番','形式: WO-YYYYMMDD-NN（例: WO-20260514-01）'),('④','作業実施','ドライバーが収集・運搬を実施'),
      ('⑤','ドライバー実績記入','帰社後、作業指示書の実績欄に記入して提出'),('⑥','事務員：売上データ入力','翌日以降、担当事務員がシステムに入力'),
      ('⑦','担当者チェック（確認済）','担当者が内容を確認し「確認済」ステータスへ変更'),('⑧','月末：CSV出力','確認済の案件を対象に月次CSVを出力'),
      ('⑨','T-PLANNERインポート','CSVをT-PLANNERに取込み'),('⑩','売上伝票計上','売上伝票としてT-PLANNERに登録完了')]
tf=doc.add_table(rows=len(flow),cols=3); tf.style='Table Grid'
for ri,(num,step,desc) in enumerate(flow):
    row=tf.rows[ri]; rb=RGBColor(0xdb,0xe4,0xf8) if ri%2==0 else RGBColor(0xef,0xf6,0xff)
    rb2=RGBColor(0xff,0xff,0xff) if ri%2==0 else RGBColor(0xf8,0xfa,0xfc)
    for ci,(v,w) in enumerate(zip([num,step,desc],[Cm(1.2),Cm(5.5),Cm(9.3)])):
        c=row.cells[ci]; c.width=w; cell_bg(c,rb if ci<2 else rb2); cell_bd(c,'BFDBFE',4)
        p=c.paragraphs[0]; p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
        if ci==0: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run(p,v,pt=10,bold=(ci<2),color=C_NAVY if ci<2 else None)
para('')
note('【重要ポイント】\n・売上計上が必要な案件と不要な案件を、配車カード作成時点で区別する。\n・作業指示書番号（WO番号）が、作業指示書と売上伝票を一対一で紐づけるキーとなる。\n・得意先CDは必ずお客様マスタから選択し、誤請求を構造的に防止する。',bc='2563EB')
doc.add_page_break()

# ═══ 第2章 ═══
h1('第2章　収集タイプと売上計上の関係')
para('収集の方法（当社車両で出向するか、お客様が持込むか）によって、計上する費目が異なります。',pt=10.5,sa=8)
table(['収集タイプ','説明','廃棄物処理費','収集運搬費','備考'],
      [['通常収集（車両出向）','当社車両でお客様先へ収集に行く','◎ 計上','◎ 計上','ドライバー配車あり'],
       ['持込み','お客様が当社施設へ持ち込む','◎ 計上','✕ 計上なし','車両費は発生しない'],
       ['計上なし','社内処理・無償対応等','─','─','売上伝票不要']],
      [Cm(3.0),Cm(4.2),Cm(2.8),Cm(2.8),Cm(3.2)])
para('')
note('【実装上の扱い】\n配車カードが「通常収集（車両出向）」の場合、売上明細入力画面で収集運搬費品目を自動的に追加します。持込みの場合は自動追加しません（手動追加は可能）。',bc='16A34A')
doc.add_page_break()

# ═══ 第3章 ═══
h1('第3章　必要マスタデータ')
para('「既存」は現在のシステムにあるもの、「★新規追加」は今回追加が必要なものです。',pt=10.5,sa=8)
h2('3-1　お客様マスタ（既存＋拡張）')
para('【重要】同一お客様が請求種別（定額・スポット等）ごとに複数の得意先CDを持つ場合があります。1社あたり複数の得意先CDを登録し、作業指示書作成時に使用するCDを選択します。',pt=10,sa=6)
table(['フィールド','内容','種別','備考'],
      [['お客様ID','システム内部ID','既存',''],['お客様名','会社名','既存',''],['フリガナ','カナ表記','既存',''],
       ['得意先CDリスト','複数登録可能','★新規追加','1社につき複数登録可'],
       ['　- 得意先CD','T-PLANNERコード','★新規追加','例：1001'],
       ['　- 用途ラベル','定額契約用／スポット用など','★新規追加','例：スポット案件用'],
       ['　- 備考','使用条件等メモ','★新規追加','任意入力']],
      [Cm(3.5),Cm(4.0),Cm(2.5),Cm(6.0)])
para('')
h2('3-2　品目マスタ（新規）')
table(['フィールド','内容','備考'],
      [['品目CD','T-PLANNER品目コード','T-PLANNERからCSV取得'],['品目名','廃棄物処理費・収集運搬費など',''],
       ['標準単位','m³・式・kgなど',''],['標準単価','デフォルト単価','顧客別単価で上書き可'],
       ['収集運搬費フラグ','この品目が車両費かどうか','自動追加判定に使用（◯/✕）']],
      [Cm(3.5),Cm(5.5),Cm(7.0)])
para('')
h2('3-3　顧客別単価マスタ（新規）')
para('単価の優先順位：① 顧客×品目の個別単価（最優先）　② 品目の標準単価　③ 手動入力による上書き',pt=10,sa=6)
table(['フィールド','内容','備考'],
      [['お客様ID','紐づくお客様',''],['品目CD','対象品目',''],
       ['単価','この顧客向け単価','円単位'],['適用開始日','単価改定時に対応','改定前の単価は履歴として保持']],
      [Cm(3.5),Cm(5.0),Cm(7.5)])
doc.add_page_break()

# ═══ 第4章 ═══
h1('第4章　作業指示書番号（採番ルール）')
para('作業指示書と売上伝票を確実に1対1で紐づけるため、ユニークな作業指示書番号を自動採番します。',pt=10.5,sa=8)
h2('4-1　採番形式')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)
run(p,'WO - YYYYMMDD - NN',pt=20,bold=True,color=C_NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
run(p,'例：WO-20260514-01　（2026年5月14日 1件目）',pt=11,color=C_BLUE)
para('')
table(['要素','内容','例'],[['WO','固定プレフィックス（Work Order）','WO'],['YYYYMMDD','配車日（8桁）','20260514'],['NN','当日の連番（2桁・01から開始）','01, 02, 03 …']],[Cm(3.5),Cm(8.5),Cm(4.0)])
para('')
h2('4-2　特徴')
for item in ['自動採番：システムが配車日と連番を組み合わせて生成（手動変更不可）',
             '配車日から案件を特定しやすい（日付が番号に含まれる）',
             'マニフェスト番号とは別管理（マニフェストが発生しない案件も存在するため）',
             'WO番号が作業指示書とT-PLANNER売上伝票を1対1で結びつけるキーとなる',
             'CSV出力後も番号は変更できない（証跡・照合のため）']:
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(2); run(p,item,pt=10)
doc.add_page_break()

# ═══ 第5章 ═══
h1('第5章　売上データ入力画面（概念設計）')
para('画面はモーダルダイアログとして配車管理システム上に表示されます。',pt=10.5,sa=8)
table(['エリア','項目名','内容・操作'],
      [['【エリア①】基本情報','作業指示書No','自動採番・表示のみ（WO-YYYYMMDD-NN）'],
       ['','配車日','配車カードから自動取得'],['','お客様','配車カードから自動取得'],
       ['','得意先CD','お客様に紐づく得意先CDをドロップダウン選択（複数登録がある場合は選択必須）'],
       ['','収集先','配車カードから自動取得'],['','ドライバー／車両','配車カードから自動取得'],['','マニフェストNo','手動入力（任意）'],
       ['【エリア②】売上計上区分','計上区分','◉ 計上あり / ○ 計上なし（社内・無償）/ ○ 保留'],
       ['【エリア③】売上明細','収集運搬費フラグ','◉ 発生あり（車両出向）→ 収集運搬費を自動追加 / ○ なし（持込み）'],
       ['','品目','プルダウンから品目CDを選択（品目名・単価を自動表示）'],
       ['','数量／単位／単価','数量と単位を入力、単価は顧客別単価マスタから自動取得（変更可）'],
       ['','金額／合計','数量×単価を自動計算'],
       ['【エリア④】ドライバー実績','収集物内容・実績数量・特記事項','自由記述（帰社後に事務員が入力）'],
       ['【エリア⑤】ステータス','ステータス','入力中 → 入力完了 → 確認済 → CSV出力済'],
       ['【エリア⑥】操作履歴','履歴ログ','作成・変更・削除の日時と操作者を自動記録']],
      [Cm(3.2),Cm(3.2),Cm(9.6)])
para('')
note('【画面下部のボタン】\n💾 保存（入力中→入力完了）　｜　✅ 確認済にする　｜　📊 CSV出力（確認済のみ）　｜　🖨 印刷（作業指示書として出力）',bc='2563EB')
doc.add_page_break()

# ═══ 第6章 ═══
h1('第6章　ステータス管理フロー')
para('売上データの入力からT-PLANNERへの反映まで、4段階のステータスで管理します。',pt=10.5,sa=8)
h2('6-1　ステータス遷移')
t6=doc.add_table(rows=1,cols=7); t6.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,(l,w,rb,tc) in enumerate(zip(
        ['入力中','▶','入力完了','▶','確認済','▶','CSV出力済'],
        [Cm(2.5),Cm(0.8),Cm(2.5),Cm(0.8),Cm(2.5),Cm(0.8),Cm(2.5)],
        [RGBColor(0xf1,0xf5,0xf9),None,RGBColor(0xdb,0xe4,0xf8),None,RGBColor(0xdc,0xfc,0xe7),None,RGBColor(0xff,0xed,0xd5)],
        [C_GRAY,C_GRAY,C_BLUE,C_GRAY,C_GREEN,C_GRAY,C_ORANGE])):
    c=t6.rows[0].cells[i]; c.width=w
    if rb: cell_bg(c,rb); cell_bd(c,'D1D5DB',4)
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(6)
    run(p,l,pt=10,bold=(rb is not None),color=tc)
para('')
h2('6-2　各ステータスの詳細')
table(['ステータス','操作者','説明','編集可否','次のアクション'],
      [['入力中','事務員','データ入力作業中','編集可','「保存」→「入力完了」へ'],
       ['入力完了','事務員','入力が終わり確認待ち','編集可（再度「入力完了」が必要）','担当者へ確認依頼'],
       ['確認済','担当者','内容チェック済・CSV出力待ち','担当者のみ変更可','月末CSV出力へ'],
       ['CSV出力済','システム','T-PLANNERへインポート済','変更不可（二重計上防止）','─']],
      [Cm(2.5),Cm(2.0),Cm(3.5),Cm(3.5),Cm(4.5)])
doc.add_page_break()

# ═══ 第7章 ═══
h1('第7章　月次CSV出力フロー（T-PLANNERインポート）')
para('月末（または任意の締め日）に、確認済ステータスの案件をまとめてCSV出力し、T-PLANNERへインポートします。',pt=10.5,sa=8)
h2('7-1　推奨運用フロー')
table(['タイミング','アクション','ポイント'],
      [['月中（随時）','作業指示書入力 → 「確認済」へ変更','この時点ではT-PLANNERに入れない'],
       ['月末・締め処理','「確認済」の件数・金額を一覧で確認','担当者が内容を最終チェック'],
       ['CSV出力','「月次CSV出力」ボタンで一括ダウンロード','対象：「確認済」ステータスのみ'],
       ['T-PLANNERインポート','ダウンロードしたCSVをT-PLANNERに取込み','1行＝1品目明細'],
       ['インポート完了','ステータスを「CSV出力済」に自動変更','以降は変更不可（二重計上防止）']],
      [Cm(3.0),Cm(7.0),Cm(6.0)])
para('')
h2('7-2　T-PLANNER CSVフォーマット対応')
table(['CSVカラム','内容','取得元'],
      [['売上日','配車日（配車確定日）','作業指示書の配車日'],
       ['得意先CD','T-PLANNER得意先コード','作業指示書で選択した得意先CD'],
       ['品目CD','T-PLANNER品目コード','売上明細の品目CD'],
       ['数量','数値','売上明細の数量'],['単位','テキスト','売上明細の単位'],
       ['単価','数値','売上明細の単価'],['金額','数値（単価×数量）','自動計算']],
      [Cm(3.0),Cm(4.0),Cm(9.0)])
doc.add_page_break()

# ═══ 第8章 ═══
h1('第8章　操作履歴の記録仕様')
para('「いつ・誰が・何を」操作したかを自動で記録します。誤入力の発見や責任の明確化に役立てます。',pt=10.5,sa=8)
h2('8-1　記録タイミング')
for item in ['新規作成時（作業指示書の初回保存）','任意フィールドの変更時（変更前・変更後の値を記録）',
             'ステータス変更時（例：入力完了 → 確認済）','削除時（論理削除・データは残す。完全削除は行わない）']:
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(2); run(p,item,pt=10)
para('')
h2('8-2　記録項目')
table(['記録項目','内容','例'],
      [['操作日時','年月日時分秒','2026-05-16 14:05:32'],['操作者','ログインスタッフ名','佐藤 一郎'],
       ['操作種別','作成 / 変更 / ステータス変更 / 削除','変更'],['変更内容','変更前・変更後の値を記録','単価: 8,000 → 9,000']],
      [Cm(3.0),Cm(6.0),Cm(7.0)])
doc.add_page_break()

# ═══ 第9章 ═══
h1('第9章　未決定事項（確認が必要な項目）')
para('以下の項目について、業務担当者・スタッフとの確認が必要です。確認後、本設計書に反映します。',pt=10.5,sa=8)
table(['No','確認事項','選択肢・補足','確認状況'],
      [['1','ステータス承認フローの深さ','現案：4段階。責任者承認ステップが別途必要か？','未確認'],
       ['2','収集運搬費の品目CD','T-PLANNERでの品目CDは何番か？（自動追加に必要）','未確認'],
       ['3','月次締め日','末日固定か？20日締めなど任意の締め日があるか？','未確認'],
       ['4','スタッフマスタの扱い','操作履歴記録のため、既存スタッフデータを使用するか？','未確認'],
       ['5','品目マスタの登録方法','T-PLANNERからCSVエクスポートして一括インポートか？','未確認'],
       ['6','顧客別単価の初期登録','T-PLANNERの単価データをCSVエクスポートできるか？','未確認'],
       ['7','CSV出力の実施頻度','月末1回か？月中に複数回出力することがあるか？','未確認']],
      [Cm(0.8),Cm(3.5),Cm(9.0),Cm(2.7)])
para('')
note('【次のアクション】\n上記の確認事項を業務担当者・スタッフと打ち合わせのうえ回答を記入し、本設計書を更新してください。確認完了後、実装フェーズに移行します。',bc='1e3a5f')

# ── ヘッダー / フッター ──
from docx.oxml import OxmlElement as oxe; from docx.oxml.ns import qn as ns2
hdr=section.header; hp=hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
hp.clear(); hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; run(hp,'配車管理システム 売上計上モジュール 設計書',pt=8,color=C_GRAY)
ftr=section.footer; fp=ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
fp.clear(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; run(fp,'- ',pt=9,color=C_GRAY)
rp=fp.add_run(); rp.font.name=FONT; rp.font.size=Pt(9)
fc1=oxe('w:fldChar'); fc1.set(ns2('w:fldCharType'),'begin')
it=oxe('w:instrText'); it.text='PAGE'; it.set('{http://www.w3.org/XML/1998/namespace}space','preserve')
fc2=oxe('w:fldChar'); fc2.set(ns2('w:fldCharType'),'end')
rp._r.append(fc1); rp._r.append(it); rp._r.append(fc2); run(fp,' -',pt=9,color=C_GRAY)

out=os.path.join(os.path.dirname(os.path.abspath(__file__)),'売上計上システム設計書.docx')
doc.save(out); print(f'Done: {out}')
